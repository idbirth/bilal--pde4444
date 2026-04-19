#!/usr/bin/env python3
from __future__ import annotations

import re
import subprocess
import sys
from zipfile import ZipFile
from pathlib import Path

_VENDOR_DIR = Path(__file__).resolve().parent / ".vendor"
if _VENDOR_DIR.exists():
    sys.path.insert(0, str(_VENDOR_DIR))

import markdown


def run(cmd: list[str]) -> None:
    subprocess.run(cmd, check=True)


def convert_template_to_html(template_docx: Path, temp_html: Path) -> str:
    run([
        "textutil",
        "-convert",
        "html",
        "-output",
        str(temp_html),
        str(template_docx),
    ])
    return temp_html.read_text(encoding="utf-8")


def extract_style_block(template_html: str) -> str:
    match = re.search(r"(<style type=\"text/css\">.*?</style>)", template_html, re.S)
    if not match:
        raise RuntimeError("Could not extract style block from template HTML.")
    return match.group(1)


def build_html_document(style_block: str, body_html: str) -> str:
    return f"""<!DOCTYPE html PUBLIC "-//W3C//DTD HTML 4.01//EN" "http://www.w3.org/TR/html4/strict.dtd">
<html>
<head>
  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
  <meta http-equiv="Content-Style-Type" content="text/css">
  <title>Claude Work Report</title>
  {style_block}
  <style type="text/css">
    body {{
      margin: 36px 54px;
      font: 12.0px Times;
      color: #111111;
    }}
    h1 {{
      margin: 0.0px 0.0px 12.0px 0.0px;
      text-align: center;
      font: 20.0px Georgia;
    }}
    h2 {{
      margin: 18.0px 0.0px 8.0px 0.0px;
      font: 17.0px Georgia;
    }}
    h3 {{
      margin: 14.0px 0.0px 6.0px 0.0px;
      font: 14.0px Georgia;
    }}
    h4 {{
      margin: 10.0px 0.0px 4.0px 0.0px;
      font: 12.0px Georgia;
    }}
    p, li {{
      font: 12.0px Times;
      line-height: 1.35;
    }}
    ul, ol {{
      margin-top: 6px;
      margin-bottom: 6px;
    }}
    table {{
      border-collapse: collapse;
      margin: 10px 0 14px 0;
      width: 100%;
    }}
    th, td {{
      border: 1px solid #bfbfbf;
      padding: 6px 7px;
      vertical-align: top;
      font: 11.0px Times;
    }}
    th {{
      background: #f3f3f3;
    }}
    img {{
      display: block;
      max-width: 100%;
      height: auto;
      margin: 12px auto 6px auto;
    }}
    blockquote {{
      margin: 10px 0;
      padding-left: 10px;
      border-left: 3px solid #bfbfbf;
      color: #333333;
    }}
    code {{
      font-family: Menlo, Consolas, monospace;
      font-size: 10.5px;
    }}
    pre {{
      border: 1px solid #d9d9d9;
      background: #fafafa;
      padding: 8px;
      overflow-x: auto;
    }}
    hr {{
      border: none;
      border-top: 1px solid #d9d9d9;
      margin: 18px 0;
    }}
    .front-matter {{
      text-align: center;
      margin-bottom: 18px;
    }}
  </style>
</head>
<body>
  <div class="front-matter">
    <p><b>Module:</b> Machine Learning for Engineers</p>
    <p><b>ASSESSMENT: ML-Based Visual Quality Inspection System</b></p>
  </div>
  {body_html}
</body>
</html>
"""


def absolutize_local_image_sources(body_html: str, base_dir: Path) -> str:
    def repl(match: re.Match[str]) -> str:
        src = match.group(1)
        if src.startswith(("http://", "https://", "file://", "data:")):
            return match.group(0)
        resolved = (base_dir / src).resolve()
        return f'src="{resolved.as_uri()}"'

    return re.sub(r'src="([^"]+)"', repl, body_html)


def count_embedded_media(docx_path: Path) -> int:
    if not docx_path.exists():
        return 0
    with ZipFile(docx_path) as archive:
        return len([name for name in archive.namelist() if name.startswith("word/media/")])


def add_text_run(paragraph, text: str | None, *, bold: bool = False, italic: bool = False,
                 code: bool = False, underline: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if code:
        run.font.name = "Courier New"


def add_inline_runs(paragraph, element, *, bold: bool = False, italic: bool = False,
                    code: bool = False, underline: bool = False) -> None:
    add_text_run(paragraph, element.text, bold=bold, italic=italic, code=code, underline=underline)
    for child in element:
        tag = child.tag.lower() if isinstance(child.tag, str) else ""
        child_bold = bold or tag in {"strong", "b"}
        child_italic = italic or tag in {"em", "i"}
        child_code = code or tag == "code"
        child_underline = underline or tag == "a"

        if tag == "br":
            paragraph.add_run().add_break()
        elif tag == "img":
            pass
        else:
            add_inline_runs(
                paragraph,
                child,
                bold=child_bold,
                italic=child_italic,
                code=child_code,
                underline=child_underline,
            )

        add_text_run(paragraph, child.tail, bold=bold, italic=italic, code=code, underline=underline)


def set_document_styles(document) -> None:
    from docx.shared import Pt

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    heading_sizes = {
        "Title": 20,
        "Heading 1": 17,
        "Heading 2": 14,
        "Heading 3": 12,
    }
    for style_name, size in heading_sizes.items():
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Georgia"
            style.font.size = Pt(size)
            style.font.bold = True


def build_docx_with_python_docx(body_html: str, output_docx: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from lxml import html as lxml_html

    def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = bold

    def shade_header_cell(cell) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F3F3F3")
        tc_pr.append(shd)

    def render_block(document, block) -> None:
        tag = block.tag.lower() if isinstance(block.tag, str) else ""

        if tag == "h1":
            document.add_heading(block.text_content().strip(), level=0)
            return
        if tag == "h2":
            document.add_heading(block.text_content().strip(), level=1)
            return
        if tag == "h3":
            document.add_heading(block.text_content().strip(), level=2)
            return
        if tag == "h4":
            document.add_heading(block.text_content().strip(), level=3)
            return

        if tag == "p":
            images = block.findall(".//img")
            if images:
                for image in images:
                    src = image.get("src")
                    if not src:
                        continue
                    if src.startswith("file://"):
                        image_path = Path(src.removeprefix("file://"))
                    else:
                        image_path = Path(src)
                    if image_path.exists():
                        document.add_picture(str(image_path), width=Inches(5.8))
                return

            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, block)
            if not paragraph.text.strip():
                paragraph._element.getparent().remove(paragraph._element)
            return

        if tag in {"ul", "ol"}:
            style = "List Bullet" if tag == "ul" else "List Number"
            for li in block.findall("./li"):
                paragraph = document.add_paragraph(style=style)
                add_inline_runs(paragraph, li)
            return

        if tag == "blockquote":
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, block, italic=True)
            return

        if tag == "table":
            rows = block.findall("./thead/tr") + block.findall("./tbody/tr")
            if not rows:
                rows = block.findall("./tr")
            if not rows:
                return

            col_count = max(
                len(row.findall("./th")) + len(row.findall("./td"))
                for row in rows
            )
            table = document.add_table(rows=len(rows), cols=col_count)
            table.style = "Table Grid"

            for row_idx, row in enumerate(rows):
                cells = row.findall("./th") + row.findall("./td")
                for col_idx, cell_elem in enumerate(cells):
                    text = " ".join(part.strip() for part in cell_elem.itertext() if part.strip())
                    is_header = cell_elem.tag.lower() == "th" or row_idx == 0
                    set_cell_text(table.cell(row_idx, col_idx), text, bold=is_header)
                    if is_header:
                        shade_header_cell(table.cell(row_idx, col_idx))
            document.add_paragraph()
            return

        if tag == "hr":
            document.add_paragraph()
            return

    document = Document()
    set_document_styles(document)

    section = document.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    front = document.add_paragraph()
    front.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = front.add_run("Module: ")
    run1.bold = True
    front.add_run("Machine Learning for Engineers")

    front2 = document.add_paragraph()
    front2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = front2.add_run("ASSESSMENT: ")
    run2.bold = True
    front2.add_run("ML-Based Visual Quality Inspection System")

    root = lxml_html.fragment_fromstring(body_html, create_parent="div")
    for child in root:
        render_block(document, child)

    document.save(str(output_docx))


def main() -> int:
    base_dir = Path(__file__).resolve().parent
    root_dir = base_dir.parent
    template_docx = root_dir / "Assessment Details UPDATED.docx"
    markdown_file = base_dir / "Claude_Work_Report.md"
    output_docx = base_dir / "Claude_Work_Report.docx"
    temp_template_html = base_dir / ".template_from_docx.html"
    temp_report_html = base_dir / ".report_build.html"

    if not template_docx.exists():
        raise FileNotFoundError(f"Template not found: {template_docx}")
    if not markdown_file.exists():
        raise FileNotFoundError(f"Markdown report not found: {markdown_file}")

    template_html = convert_template_to_html(template_docx, temp_template_html)
    style_block = extract_style_block(template_html)

    md_text = markdown_file.read_text(encoding="utf-8")
    body_html = markdown.markdown(
        md_text,
        extensions=[
            "tables",
            "fenced_code",
            "sane_lists",
        ],
        output_format="html5",
    )
    body_html = absolutize_local_image_sources(body_html, base_dir)

    final_html = build_html_document(style_block, body_html)
    temp_report_html.write_text(final_html, encoding="utf-8")

    run([
        "textutil",
        "-convert",
        "docx",
        "-baseurl",
        base_dir.resolve().as_uri() + "/",
        "-output",
        str(output_docx),
        str(temp_report_html),
    ])

    if count_embedded_media(output_docx) == 0 and 'src="file://' in body_html:
        build_docx_with_python_docx(body_html, output_docx)

    temp_template_html.unlink(missing_ok=True)
    temp_report_html.unlink(missing_ok=True)
    print(f"Built: {output_docx}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"build_report.py failed: {exc}", file=sys.stderr)
        raise
